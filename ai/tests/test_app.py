import docx


def test_health(client):
    response = client.get("/health")
    assert response.status_code == 200
    assert response.get_json()["status"] == "ok"


def test_match(client):
    payload = {
        "resume_text": "Python developer with SQL and Django experience.",
        "job": {
            "title": "Backend Engineer",
            "description": "Looking for Django and SQL skills.",
            "required_skills": ["python", "django", "sql"]
        }
    }
    response = client.post("/match", json=payload)
    assert response.status_code == 200
    data = response.get_json()
    assert 0 <= data["score"] <= 100
    assert "python" in data["matched_skills"]


def test_recommend_jobs(client):
    payload = {
        "resume_text": "Data analyst with Python and SQL.",
        "jobs": [
            {"id": 1, "title": "Data Analyst", "description": "SQL and Python required", "required_skills": ["python", "sql"]},
            {"id": 2, "title": "Designer", "description": "Figma and UX", "required_skills": ["figma"]}
        ],
        "top_n": 1
    }
    response = client.post("/recommend_jobs", json=payload)
    assert response.status_code == 200
    data = response.get_json()
    assert len(data["recommendations"]) == 1
    assert data["recommendations"][0]["job_id"] == 1


def test_skill_gap(client):
    payload = {
        "resume_text": "Python developer with SQL.",
        "job": {
            "title": "Backend Engineer",
            "description": "Looking for Django and SQL skills.",
            "required_skills": ["python", "django", "sql"]
        }
    }
    response = client.post("/skill_gap", json=payload)
    assert response.status_code == 200
    data = response.get_json()
    assert "django" in data["missing_skills"]


def test_parse_resume_docx(client, tmp_path):
    doc = docx.Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Email: jane@example.com")
    doc.add_paragraph("Skills: Python, SQL, Django")
    file_path = tmp_path / "resume.docx"
    doc.save(file_path)

    response = client.post("/parse_resume", json={
        "file_path": str(file_path),
        "file_type": "docx"
    })
    assert response.status_code == 200
    data = response.get_json()
    assert "python" in data["skills"]
